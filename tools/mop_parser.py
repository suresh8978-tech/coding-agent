"""MOP (Method of Procedure) document parser for the coding agent."""

from pathlib import Path
from typing import Any
from langchain_core.tools import tool
from tools.utils import safe_tool


@tool
@safe_tool
def read_mop_document(path: str) -> dict[str, Any]:
    """Read and parse a MOP document (DOCX format, up to 90 pages).
    
    This tool extracts all text content, tables, headings, and structure
    from a Microsoft Word document to enable the agent to understand
    and execute the procedures described within.
    
    Args:
        path: Path to the DOCX file.
        
    Returns:
        A dictionary containing:
        - title: Document title (if found)
        - paragraphs: List of paragraph texts
        - headings: List of headings with their level
        - tables: List of tables (each as list of rows)
        - full_text: Complete document text
        - metadata: Document metadata (author, created, etc.)
    """
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    
    file_path = Path(path)
    if not file_path.exists():
        return {"error": f"File '{path}' does not exist."}
    if not file_path.suffix.lower() == '.docx':
        return {"error": f"File '{path}' is not a DOCX file."}
    
    try:
        doc = Document(str(file_path))
    except PackageNotFoundError:
        return {"error": f"Cannot open '{path}'. File may be corrupted or not a valid DOCX."}
    except Exception as e:
        return {"error": f"Error opening document: {str(e)}"}
    
    result = {
        "title": None,
        "paragraphs": [],
        "headings": [],
        "tables": [],
        "lists": [],
        "full_text": "",
        "metadata": {},
        "sections": [],
    }
    
    # Extract metadata
    try:
        core_props = doc.core_properties
        result["metadata"] = {
            "author": core_props.author,
            "title": core_props.title,
            "subject": core_props.subject,
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
        }
        if core_props.title:
            result["title"] = core_props.title
    except Exception:
        pass  # Metadata extraction is optional
    
    # Extract paragraphs and identify headings
    full_text_parts = []
    current_section = {"heading": None, "level": 0, "content": []}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if it's a heading
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if 'heading' in style_name:
                # Extract heading level
                level = 1
                for char in style_name:
                    if char.isdigit():
                        level = int(char)
                        break
                
                result["headings"].append({
                    "text": text,
                    "level": level,
                })
                
                # Save previous section and start new one
                if current_section["heading"] or current_section["content"]:
                    result["sections"].append(current_section)
                current_section = {"heading": text, "level": level, "content": []}
                
                # Set document title from first heading if not already set
                if result["title"] is None and level == 1:
                    result["title"] = text
            else:
                result["paragraphs"].append(text)
                current_section["content"].append(text)
        else:
            result["paragraphs"].append(text)
            current_section["content"].append(text)
        
        full_text_parts.append(text)
    
    # Save last section
    if current_section["heading"] or current_section["content"]:
        result["sections"].append(current_section)
    
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        
        if table_data:
            result["tables"].append(table_data)
            # Also add table content to full text
            for row in table_data:
                full_text_parts.append(" | ".join(row))
    
    result["full_text"] = "\n".join(full_text_parts)
    
    # Summary statistics
    result["stats"] = {
        "paragraph_count": len(result["paragraphs"]),
        "heading_count": len(result["headings"]),
        "table_count": len(result["tables"]),
        "section_count": len(result["sections"]),
        "character_count": len(result["full_text"]),
        "word_count": len(result["full_text"].split()),
    }
    
    return result
