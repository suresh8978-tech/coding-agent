"""Document creation tools for the coding agent.

Supports creating documents in DOCX, PDF, TXT, and Markdown formats
from Markdown-formatted content.
"""

import os
import re
from pathlib import Path
from langchain_core.tools import tool
from tools.utils import safe_tool


def _parse_markdown_content(content: str) -> list[dict]:
    """Parse Markdown content into structured blocks.

    Returns a list of dicts with keys:
        - type: 'heading', 'paragraph', 'bullet', 'table'
        - level: heading level (1-6) for headings
        - text: the text content
        - rows: list of lists for tables
    """
    blocks = []
    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Heading
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append({"type": "heading", "level": level, "text": heading_match.group(2).strip()})
            i += 1
            continue

        # Bullet list item
        bullet_match = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet_match:
            blocks.append({"type": "bullet", "text": bullet_match.group(1).strip()})
            i += 1
            continue

        # Numbered list item
        numbered_match = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if numbered_match:
            blocks.append({"type": "numbered", "text": numbered_match.group(1).strip()})
            i += 1
            continue

        # Table (pipe-delimited)
        if "|" in line and line.strip().startswith("|"):
            table_rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                # Skip separator lines like |---|---|
                if re.match(r"^\|[\s\-:|]+\|$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")[1:-1]]
                table_rows.append(cells)
                i += 1
            if table_rows:
                blocks.append({"type": "table", "rows": table_rows})
            continue

        # Non-empty paragraph line
        if line.strip():
            # Collect consecutive non-empty, non-special lines as one paragraph
            para_lines = []
            while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6}\s|[-*+]\s|\d+[.)]\s|\|)", lines[i]):
                para_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
            continue

        # Empty line — skip
        i += 1

    return blocks


def _create_docx(file_path: str, title: str, blocks: list[dict]) -> str:
    """Create a DOCX document from parsed blocks."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for block in blocks:
        if block["type"] == "heading":
            doc.add_heading(block["text"], level=min(block["level"], 9))

        elif block["type"] == "paragraph":
            doc.add_paragraph(block["text"])

        elif block["type"] == "bullet":
            doc.add_paragraph(block["text"], style="List Bullet")

        elif block["type"] == "numbered":
            doc.add_paragraph(block["text"], style="List Number")

        elif block["type"] == "table":
            rows = block["rows"]
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < num_cols:
                        table.rows[r_idx].cells[c_idx].text = cell_text

    doc.save(file_path)
    return f"Successfully created DOCX document at '{file_path}'."


def _create_pdf(file_path: str, title: str, blocks: list[dict]) -> str:
    """Create a PDF document from parsed blocks."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 20)
    pdf.multi_cell(0, 15, title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    for block in blocks:
        if block["type"] == "heading":
            level = block["level"]
            size = max(18 - (level - 1) * 2, 10)
            pdf.set_font("Helvetica", "B", size)
            pdf.ln(3)
            pdf.multi_cell(0, 7, block["text"], new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

        elif block["type"] == "paragraph":
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, block["text"], new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

        elif block["type"] == "bullet":
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, "  - " + block["text"], new_x="LMARGIN", new_y="NEXT")

        elif block["type"] == "numbered":
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, "  - " + block["text"], new_x="LMARGIN", new_y="NEXT")

        elif block["type"] == "table":
            rows = block["rows"]
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            page_width = pdf.w - pdf.l_margin - pdf.r_margin
            col_width = page_width / num_cols

            for r_idx, row in enumerate(rows):
                pdf.set_font("Helvetica", "B" if r_idx == 0 else "", 10)
                for c_idx in range(num_cols):
                    cell_text = row[c_idx] if c_idx < len(row) else ""
                    pdf.cell(col_width, 8, cell_text, border=1)
                pdf.ln()
            pdf.ln(3)

    pdf.output(file_path)
    return f"Successfully created PDF document at '{file_path}'."


def _create_txt(file_path: str, title: str, blocks: list[dict]) -> str:
    """Create a plain-text document from parsed blocks."""
    lines = []
    lines.append(title.upper())
    lines.append("=" * len(title))
    lines.append("")

    for block in blocks:
        if block["type"] == "heading":
            lines.append("")
            lines.append(block["text"].upper() if block["level"] <= 2 else block["text"])
            char = "=" if block["level"] == 1 else "-"
            lines.append(char * len(block["text"]))
            lines.append("")

        elif block["type"] == "paragraph":
            lines.append(block["text"])
            lines.append("")

        elif block["type"] == "bullet":
            lines.append(f"  * {block['text']}")

        elif block["type"] == "numbered":
            lines.append(f"  - {block['text']}")

        elif block["type"] == "table":
            rows = block["rows"]
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            col_widths = [0] * num_cols
            for row in rows:
                for c_idx, cell in enumerate(row):
                    if c_idx < num_cols:
                        col_widths[c_idx] = max(col_widths[c_idx], len(cell))
            for r_idx, row in enumerate(rows):
                padded = []
                for c_idx in range(num_cols):
                    cell = row[c_idx] if c_idx < len(row) else ""
                    padded.append(cell.ljust(col_widths[c_idx]))
                lines.append("| " + " | ".join(padded) + " |")
                if r_idx == 0:
                    lines.append("|-" + "-|-".join("-" * w for w in col_widths) + "-|")
            lines.append("")

    Path(file_path).parent.mkdir(parents=True, exist_ok=True)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return f"Successfully created TXT document at '{file_path}'."


def _create_md(file_path: str, title: str, content: str) -> str:
    """Create a Markdown document (adds title heading + raw content)."""
    md_content = f"# {title}\n\n{content}"

    Path(file_path).parent.mkdir(parents=True, exist_ok=True)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    return f"Successfully created Markdown document at '{file_path}'."


# Maps extensions to format keys
_EXT_TO_FORMAT = {
    ".docx": "docx",
    ".doc": "docx",
    ".pdf": "pdf",
    ".txt": "txt",
    ".text": "txt",
    ".md": "md",
    ".markdown": "md",
}


@tool
@safe_tool
def create_document(
    file_path: str,
    title: str,
    content: str,
    format: str = "",
) -> str:
    """Create a professional document in DOCX, PDF, TXT, or Markdown format.

    The content should be provided in Markdown format. The tool will parse
    headings (# / ## / ###), bullet lists (- item), numbered lists (1. item),
    pipe-delimited tables, and plain paragraphs, then render them in the
    target document format.

    Args:
        file_path: Where to save the document (absolute or relative path).
        title: The document title.
        content: The body content written in Markdown format.
        format: Output format — one of 'docx', 'pdf', 'txt', 'md'.
                If omitted, the format is auto-detected from the file extension.

    Returns:
        A success message with the output path, or an error description.
    """
    # Determine format
    fmt = format.lower().strip() if format else ""
    if not fmt:
        ext = Path(file_path).suffix.lower()
        fmt = _EXT_TO_FORMAT.get(ext, "")
    if not fmt:
        return (
            f"Error: Cannot determine document format from path '{file_path}'. "
            "Please specify the 'format' argument as one of: docx, pdf, txt, md."
        )


    # Ensure parent directory exists
    Path(file_path).parent.mkdir(parents=True, exist_ok=True)

    # Parse content into structured blocks
    blocks = _parse_markdown_content(content)

    if fmt == "docx":
        return _create_docx(file_path, title, blocks)
    elif fmt == "pdf":
        return _create_pdf(file_path, title, blocks)
    elif fmt == "txt":
        return _create_txt(file_path, title, blocks)
    elif fmt == "md":
        return _create_md(file_path, title, content)
    else:
        return (
            f"Error: Unsupported format '{fmt}'. "
            "Supported formats are: docx, pdf, txt, md."
        )
